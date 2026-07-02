# App.py
import streamlit as st
import requests
import io
import pandas as pd
from docx import Document
from docx import Document as ReadDocument
import json
from datetime import datetime

# Set Page Config for mobile
st.set_page_config(page_title="Financial Timeline Engine", layout="centered")

# Universal Model Configuration
PRIMARY_MODEL = "google"
SECONDARY_MODEL = "groq"
TERTIARY_MODEL = "openrouter"


# Initialize session states
if "ai_connected" not in st.session_state:
    st.session_state["ai_connected"] = False

if "timeline_data" not in st.session_state:
    st.session_state["timeline_data"] = []

# File extraction layer
@st.cache_data(show_spinner=False)
def extract_document_data(uploaded_file):
    """Reads text lines from uploaded files safely."""
    if uploaded_file is None:
        return ""
    try:
        filename = uploaded_file.name.lower()
        if filename.endswith(".txt") or filename.endswith(".csv"):
            return uploaded_file.read().decode("utf-8")
        elif filename.endswith(".xlsx"):
            # EXCEL PARSER
            df_sheets = pd.read_excel(uploaded_file, sheet_name=None)
            excel_text = ""
            for sheet, df in df_sheets.items():
                excel_text += f"\n--- Excel Sheet: {sheet} ---\n" + df.to_string() + "\n"
            return excel_text
        elif filename.endswith(".docx"):
            # WORD PARSER
            word_doc = ReadDocument(uploaded_file)
            word_text = f"\n--- Word Document: {filename} ---\n"
            for para in word_doc.paragraphs:
                if para.text.strip():
                    word_text += para.text + "\n"
            return word_text
        else:
            # Fallback simple reader for byte streams
            return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        return f"Error reading file text content: {str(e)}"

# Secure AI thesis engine
import google.generativeai as genai

def call_google_ai_studio(prompt_text):
    try:
        api_key = st.secrets.get("GOOGLE_API_KEY", "")
        if not api_key: raise ValueError("Missing Google Key")
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.5-flash")
        res = model.generate_content(str(prompt_text))
        if res.text: return res.text
        raise RuntimeError("Empty response")
    except Exception: raise

def call_groq_engine(prompt_text):
    try:
        api_key = st.secrets.get("GROQ_API_KEY", "")
        if not api_key: raise ValueError("Missing Groq Key")
        endpoint = "https://groq.com"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {"model": "llama3-8b-8192", "messages": [{"role": "user", "content": str(prompt_text)}]}
        res = requests.post(endpoint, headers=headers, json=payload, timeout=30)
        if res.status_code == 200:
            return res.json()["choices"]["message"]["content"]
        raise RuntimeError(f"Failed with status: {res.status_code}")
    except Exception: raise
        
def call_openrouter_engine(prompt_text):
    """Sends financial data requests to OpenRouter securely with hard timeout retries."""
    api_key = st.secrets.get("OPENROUTER_API_KEY", "")
    if not api_key:
        return "❌ OpenRouter API Key missing inside Streamlit Secrets panel."
    
    endpoint = "https://openrouter.ai/api/v1/chat/completions"
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://streamlit.app",
        "X-Title": "Financial Timeline Engine"
    }
    
    payload = {
        "model": PRIMARY_MODEL,
        "messages": [
            {"role": "system", "content": "You are an elite Wall Street financial research analyst. Generate structured multi-section corporate reports with key dates, events, and milestones."},
            {"role": "user", "content": str(prompt_text)}
        ]
    }
    
    # Pass 1: Try Primary Model
    try:
        res = requests.post(endpoint, headers=headers, json=payload, timeout=45)
        if res.status_code == 200:
            try:
                data = res.json()
                if "choices" in data and len(data["choices"]) > 0:
                    st.session_state["ai_connected"] = True
                    return data["choices"][0]["message"]["content"]
                else:
                    return "⚠️ OpenRouter returned an empty choices payload. Please try clicking the button again."
            except Exception:
                return "⚠️ OpenRouter server returned a malformed response. The free pool is heavily congested right now. Please try again in 10 seconds!"
        else:
            return f"❌ OpenRouter Connection Failed. Server status code: {res.status_code}. Please retry."
    except requests.exceptions.Timeout:
        pass # Gracefully fall through to retry block below

    # Pass 2: Fallback to the Smart Router Net
    try:
        payload["model"] = FALLBACK_MODEL
        res = requests.post(endpoint, headers=headers, json=payload, timeout=45)
        if res.status_code == 200:
            try:
                data = res.json()
                if "choices" in data and len(data["choices"]) > 0:
                    st.session_state["ai_connected"] = True
                    return data["choices"][0]["message"]["content"]
                else:
                    return "⚠️ OpenRouter returned an empty choices payload. Please try clicking the button again."
            except Exception:
                return "⚠️ OpenRouter server returned a malformed response. The free pool is heavily congested right now. Please try again in 10 seconds!"
        else:
            return f"❌ OpenRouter Connection Failed. Server status code: {res.status_code}. Please retry."
    except Exception:
        return "🔴 AI server busy or experiencing high latency volume right now. Please tap regenerate to claim a fresh server slot link."
    
    return "⚠️ Primary AI endpoint returned an unusual response. Please check your token quota limit logs."

# Timeline extraction & parsing engine
def extract_timeline_events(ai_narrative):
    """Parses AI narrative to extract structured timeline events."""
    try:
        # Send narrative back to AI to structure it
        api_key = st.secrets.get("OPENROUTER_API_KEY", "")
        if not api_key:
            return []
        
        endpoint = "https://openrouter.ai/api/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://streamlit.app",
            "X-Title": "Financial Timeline Engine"
        }
        
        structuring_prompt = f"""Extract timeline events from this narrative and return as JSON array with objects containing: date (YYYY-MM-DD or YYYY-MM or YYYY), event (string), category (string), impact (string).

Narrative:
{ai_narrative}

Return ONLY valid JSON array, no markdown, no extra text."""
        
        payload = {
            "model": PRIMARY_MODEL,
            "messages": [{"role": "user", "content": structuring_prompt}],
            "temperature": 0.3
        }
        
        res = requests.post(endpoint, headers=headers, json=payload, timeout=45)
        if res.status_code == 200:
            data = res.json()
            if "choices" in data and len(data["choices"]) > 0:
                response_text = data["choices"][0]["message"]["content"]
                # Try to parse JSON
                try:
                    events = json.loads(response_text)
                    return events if isinstance(events, list) else []
                except:
                    return []
        return []
    except Exception:
        return []

# Micro-utility document exporter (.DOCX exporter)
def generate_docx_download(text_content, timeline_data=None):
    """Compiles the generated AI analysis report into a clean Word document download stream."""
    doc = Document()
    
    doc.add_heading("Institutional Investment Research Memo", level=1)
    doc.add_paragraph("-" * 40)
    doc.add_heading("Executive Summary & Analysis", level=2)
    
    # Secure row cleaning loop to bypass oxml crashes
    if text_content:
        clean_text_string = str(text_content)
        for line in clean_text_string.split('\n'):
            if line.strip():
                # Strip out invalid control characters safely
                sanitized_line = "".join(c for c in line if c.isprintable() or c in ['\t', '\n'])
                # Remove markdown formatting symbols
                sanitized_line = sanitized_line.replace('**', '').replace('__', '').replace('```', '')
                if sanitized_line.strip():
                    doc.add_paragraph(sanitized_line.strip())
        else:
            doc.add_paragraph("No report content generated.")
    
    # Add timeline section if data exists
    if timeline_data and len(timeline_data) > 0:
        doc.add_heading("Extracted Timeline Events", level=2)
        for event in timeline_data:
            date_str = event.get("date", "N/A")
            event_name = event.get("event", "N/A")
            category = event.get("category", "N/A")
            impact = event.get("impact", "N/A")
            
            # Sanitize timeline event strings
            date_str = "".join(c for c in str(date_str) if c.isprintable())
            event_name = "".join(c for c in str(event_name) if c.isprintable())
            category = "".join(c for c in str(category) if c.isprintable())
            impact = "".join(c for c in str(impact) if c.isprintable())
            
            doc.add_paragraph(f"📅 {date_str}: {event_name}", style="List Bullet")
            doc.add_paragraph(f"Category: {category} | Impact: {impact}", style="List Bullet 2")
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Timeline visualization engine
def render_timeline_visualization(timeline_data):
    """Renders a simplified timeline visualization for mobile."""
    if not timeline_data or len(timeline_data) == 0:
        st.info("No timeline events extracted yet.")
        return
    
    st.subheader("📊 Timeline Events")
    
    # Create a dataframe for display
    df_timeline = pd.DataFrame(timeline_data)
    
    # Display as table
    st.dataframe(df_timeline, use_container_width=True, hide_index=True)

# Main workspace control layer
def main():
    st.title("📈 Financial Timeline Engine")
    
    # Dynamic status tracker logic
    api_key_check = st.secrets.get("OPENROUTER_API_KEY", "")
    if not api_key_check:
        st.error("🔴 AI Status: Offline (Missing OpenRouter Secrets Key Mapping)")
    elif st.session_state["ai_connected"]:
        st.success("🟢 AI Status: Connected & Verified Live")
    else:
        st.info("🟡 AI Status: API Key Loaded (Awaiting First Live Document Generation Connection)")
    
    # Sidebar Document Ingestion
    st.sidebar.header("📁 Document Ingestion")
    uploaded_files = st.sidebar.file_uploader(
        "Upload Financial Documents (.txt, .csv, .xlsx, .docx)", 
        type=["txt", "csv", "xlsx", "docx"], 
        accept_multiple_files=True
    )
    
    combined_raw_text = ""
    if uploaded_files:
        for f in uploaded_files:
            combined_raw_text += f"\n--- Start of File: {f.name} ---\n"
            combined_raw_text += extract_document_data(f)
    
    # Clean executive metric data grid view summary
    st.subheader("📊 Ingested Data Summary")
    col1, col2 = st.columns(2)
    col1.metric(label="📄 Files Processed", value=len(uploaded_files))
    col2.metric(label="📊 Extracted Characters", value=len(combined_raw_text))
    
    # Trigger Action Analysis Button Link
    st.markdown("---")
    st.subheader("🔬 AI Analysis Engine")
    
    if st.button("🚀 Generate Timeline Report"):
        with st.spinner("Processing document data and generating timeline..."):
            prompt = f"""Analyze the following corporate document data text carefully. Extract key event milestones, timelines, and potential controversy flags. Write a comprehensive multi-paragraph investment memo that identifies:
1. Key financial events and dates
2. Market movements and impacts
3. Risk factors and opportunities
4. Strategic implications

Document Data:
{combined_raw_text}

Generate a professional investment memo."""
            
            ai_narrative_result = call_openrouter_engine(prompt)
            
            # Show AI Result
            st.markdown("### 📝 Generated Investment Memo")
            st.write(ai_narrative_result)
            
            # Extract timeline events
            with st.spinner("Extracting timeline events..."):
                timeline_events = extract_timeline_events(ai_narrative_result)
                st.session_state["timeline_data"] = timeline_events
            
            # Render timeline visualization
            if timeline_events:
                render_timeline_visualization(timeline_events)
            
            # Render Working Document Exporter Module Download Button Link
            if "❌" not in ai_narrative_result and "🔴" not in ai_narrative_result:
                docx_file_stream = generate_docx_download(ai_narrative_result, timeline_events)
                st.download_button(
                    label="📥 Download as Word Document",
                    data=docx_file_stream,
                    file_name="Financial_Timeline_Investment_Memo.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("Please upload financial documents to generate a report.")

def check_login():
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False
    if not st.session_state["authenticated"]:
        st.markdown("🔐 Institutional Terminal Access", unsafe_allow_html=True)
        col_l1, col_l2, col_l3 = st.columns([1, 2, 1])
        with col_l2:
            input_user = st.text_input("Username")
            input_pass = st.text_input("Password", type="password")
            if st.button("🚀 Log In", use_container_width=True):
                if input_user == "admin" and input_pass == "financial_terminal_2026":
                    st.session_state["authenticated"] = True
                    st.rerun()
                else:
                    st.error("❌ Invalid Credentials")
        return False
    return True

if __name__ == "__main__":
    if check_login():
        main()
